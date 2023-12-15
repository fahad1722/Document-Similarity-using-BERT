import time
from PyPDF2 import PdfReader
import PyPDF2
from xhtml2pdf import pisa
from django.views import View
from django.template.loader import get_template
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from django.template.loader import render_to_string
from django.http import HttpResponse
from summa import summarizer
from sklearn.decomposition import PCA
import plotly.graph_objects as go
from wordcloud import WordCloud
from io import BytesIO
import os
import tempfile
from django.conf import settings
from docx import Document
from pdfminer.high_level import extract_text_to_fp
from django.shortcuts import redirect, render
from django.http import HttpResponseRedirect
from django.core.files.storage import FileSystemStorage
from transformers import BertTokenizer, BertModel
import torch
import numpy as np
import seaborn as sns
import matplotlib.pyplot as plt
import nltk
from reportlab.lib.pagesizes import letter
from transformers import pipeline
import docx
import plotly.graph_objects as go
from docx.shared import Inches


pretrained_model_path = "C:/Users/MOHAMMED ADNAN/OneDrive/Desktop/project/Final/Transformers"
tokenizer = BertTokenizer.from_pretrained(pretrained_model_path)
model = BertModel.from_pretrained(pretrained_model_path)
device = torch.device('cuda' if torch.cuda.is_available() else 'cpu')
model.to(device)

summarization_pipeline = pipeline(
    "summarization", model="t5-base", tokenizer="t5-base")


def find_similar_sentences(doc1_text, doc2_text):
    # Split documents into sentences
    doc1_sentences = doc1_text.split('. ')
    doc2_sentences = doc2_text.split('. ')

    # Initialize similarity matrix
    similarity_matrix = np.zeros((len(doc1_sentences), len(doc2_sentences)))

    # Calculate similarity for each sentence pair
    for i, sentence1 in enumerate(doc1_sentences):
        for j, sentence2 in enumerate(doc2_sentences):
            similarity_score = calculate_similarity(sentence1, sentence2)
            similarity_matrix[i, j] = similarity_score

    # Find similar sentences based on similarity matrix
    similar_sentences = []
    for i, sentence1 in enumerate(doc1_sentences):
        for j, sentence2 in enumerate(doc2_sentences):
            if similarity_matrix[i, j] > 0.8:  # Adjust the threshold as needed
                similar_sentences.append((sentence1, sentence2))

    return similar_sentences


def create_similarity_bar_chart(similarity_scores):
    # Prepare data for the bar chart
    document_names = list(similarity_scores.keys())
    scores = list(similarity_scores.values())

    # Create the bar chart
    fig = go.Figure(data=go.Bar(x=document_names, y=scores))

    # Customize the layout
    fig.update_layout(
        title="Similarity Scores between Documents",
        xaxis_title="Documents",
        yaxis_title="Similarity Score",
        barmode="group",
        hovermode="x",
    )

    # Return the HTML representation of the chart
    return fig.to_html(full_html=False)


def generate_word_cloud(document_text):
    # Generate word cloud
    wordcloud = WordCloud(width=800, height=400,
                          background_color='white').generate(document_text)

    # Plot the word cloud
    plt.figure(figsize=(10, 6))
    plt.imshow(wordcloud, interpolation='bilinear')
    plt.axis('off')
    plt.tight_layout()
    plt.show()


def preprocess_text(text):
    tokens = tokenizer.tokenize(text)
    tokens = ['[CLS]'] + tokens[:510] + ['[SEP]']
    input_ids = tokenizer.convert_tokens_to_ids(tokens)
    input_ids = torch.tensor([input_ids]).to(device)
    return input_ids


def calculate_similarity(doc1, doc2):
    # Preprocess documents
    doc1_ids = preprocess_text(doc1)
    doc2_ids = preprocess_text(doc2)

    # Generate BERT embeddings
    with torch.no_grad():
        doc1_embeddings = model(doc1_ids)[0][:, 0, :].cpu().numpy()
        doc2_embeddings = model(doc2_ids)[0][:, 0, :].cpu().numpy()

    # Calculate cosine similarity
    similarity = np.dot(doc1_embeddings, doc2_embeddings.T) / (
        np.linalg.norm(doc1_embeddings) * np.linalg.norm(doc2_embeddings)
    )

    return similarity.item()



def document_similarity(request):
    if request.method == 'POST' and 'document1' in request.FILES and 'document2' in request.FILES:
        # Get uploaded files
        document1 = request.FILES['document1']
        document2 = request.FILES['document2']

        # Process .docx files
        if document1.name.endswith('.docx') and document2.name.endswith('.docx'):
            doc1_text = extract_text_from_docx(document1)
            doc2_text = extract_text_from_docx(document2)

        # Process .pdf files
        elif document1.name.endswith('.pdf') and document2.name.endswith('.pdf'):
            doc1_text = extract_text_from_pdf(BytesIO(document1.read()))
            doc2_text = extract_text_from_pdf(BytesIO(document2.read()))

        else:
            return render(request, 'similarity_app/error.html', {'error_message': 'Invalid file format.'})

        # Calculate document similarity
        similarity_score = calculate_similarity(doc1_text, doc2_text)

        # Generate summaries using BERT
        doc1_summary = generate_summary(doc1_text)
        doc2_summary = generate_summary(doc2_text)

        # Generate the interactive bar chart
        speedometer_chart = create_speedometer_chart(similarity_score)

        # Store the values in the session
        request.session['similarity_score'] = similarity_score
        request.session['doc1_summary'] = doc1_summary
        request.session['doc2_summary'] = doc2_summary
        request.session['speedometer_chart'] = speedometer_chart
        return render(
            request,
            'similarity_app/result.html',
            {
                'similarity_score': similarity_score,
                'speedometer_chart': speedometer_chart,
                'doc1_summary': doc1_summary,
                'doc2_summary': doc2_summary,
            }
        )

    return render(request, 'similarity_app/upload.html')


def download_report(request):
   # Retrieve the values from the session
    similarity_score = request.session.get('similarity_score')
    doc1_summary = request.session.get('doc1_summary')
    doc2_summary = request.session.get('doc2_summary')

    # Create a new Word document
    document = Document()

    # Add the report content
    document.add_heading('Document Similarity Report', level=1)
    document.add_paragraph(f'Similarity Score: {similarity_score}')

    document.add_heading('Summary', level=2)
    document.add_paragraph(f'Document 1: {doc1_summary}')
    document.add_paragraph(f'Document 2: {doc2_summary}')

    # Create a response object with the appropriate MIME type
    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

    # Set the filename for download
    response['Content-Disposition'] = 'attachment; filename="document_similarity_report.docx"'

    # Save the document to the response object
    document.save(response)

    return response


def result(request):
    # Retrieve the values from the session
    similarity_score = request.session.get('similarity_score')
    doc1_summary = request.session.get('doc1_summary')
    doc2_summary = request.session.get('doc2_summary')
    speedometer_chart = request.session.get('speedometer_chart')
    # Pass the values to the template for rendering

    return render(request, 'result.html', {
        'similarity_score': similarity_score,
        'speedometer_chart': speedometer_chart,
        'doc1_summary': doc1_summary,
        'doc2_summary': doc2_summary,
    })


def extract_text_from_docx(document):
    doc = Document(document)
    paragraphs = [p.text for p in doc.paragraphs]
    return '\n'.join(paragraphs)


def extract_text_from_pdf(document):
    pdf_reader = PdfReader(document)
    text = ''
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


def create_gauge(similarity_score):
    fig, ax = plt.subplots(subplot_kw={'aspect': 1})
    ax.set_xlim(-1, 1)
    ax.set_ylim(-1, 1)
    ax.axis('off')
    ax.add_artist(plt.Circle((0, 0), 1, color='lightgray'))
    angle = similarity_score * 180 - 90
    ax.add_artist(plt.Arrow(0, 0, 0.6 * np.cos(np.radians(angle)),
                  0.6 * np.sin(np.radians(angle)), width=0.1, color='red'))
    ax.text(
        0, -0.2, f"Similarity Score: {similarity_score:.2f}", ha='center', fontsize=12)
    return fig


def create_speedometer_chart(similarity_score):
    # Define the ranges for the speedometer chart
    ranges = [
        {'color': "green", 'start': 0, 'end': 0.4},
        {'color': "yellow", 'start': 0.4, 'end': 0.7},
        {'color': "red", 'start': 0.7, 'end': 1.0},
    ]

    # Calculate the value position on the chart
    value_position = similarity_score * 100

    # Create the speedometer chart
    fig = go.Figure()

    for rng in ranges:
        fig.add_trace(go.Indicator(
            mode="gauge+number",
            value=value_position,
            title={'text': "Similarity Score"},
            gauge={
                'axis': {'range': [0, 100]},
                'bar': {'color': rng['color']},
                'steps': [
                    {'range': [rng['start'] * 100, rng['end']
                               * 100], 'color': rng['color']}
                ],
            }
        ))

    # Customize the layout
    fig.update_layout(
        title="Document Similarity",
        width=600,
        height=400
    )

    # Return the HTML representation of the chart
    return fig.to_html(full_html=False)


def generate_summary(document_text):
    # Determine the max length based on the size of the input document
    doc_length = len(document_text)
    max_length = doc_length // 10  # Adjust this factor based on your preference

    # Use the T5 pipeline for summarization
    result = summarization_pipeline(
        document_text, max_length=max_length, min_length=30, do_sample=True, early_stopping=True)

    # Extract the generated summary from the pipeline result
    summary = result[0]['summary_text']

    return summary


class GenerateReportPDFView(View):
    def post(self, request):
        # Retrieve the necessary data for the report
        similarity_score = request.session.get('similarity_score')
        doc1_summary = request.session.get('doc1_summary')
        doc2_summary = request.session.get('doc2_summary')

        # Generate the HTML content for the report
        template = get_template('report.html')
        context = {
            'similarity_score': similarity_score,
            'doc1_summary': doc1_summary,
            'doc2_summary': doc2_summary,
        }
        html_content = template.render(context)

        # Create a PDF file using the HTML content
        result_file = 'report.pdf'
        pdf = self.generate_pdf(html_content)

        # Set the appropriate response headers for file download
        response = HttpResponse(content_type='application/pdf')
        response['Content-Disposition'] = f'attachment; filename="{result_file}"'

        # Write the PDF content to the response
        response.write(pdf)
        return response

    @staticmethod
    def generate_pdf(html_content):
        pdf_data = BytesIO()
        pisa_status = pisa.CreatePDF(html_content, dest=pdf_data)

        if pisa_status.err:
            raise Exception('PDF generation error')

        pdf_data.seek(0)
        pdf = pdf_data.read()
        pdf_data.close()
        return pdf


def show_report(request):
    # Retrieve the necessary data to generate the report
    # Retrieve the values from the session
    similarity_score = request.session.get('similarity_score')
    doc1_summary = request.session.get('doc1_summary')
    doc2_summary = request.session.get('doc2_summary')
    speedometer_chart = request.session.get('speedometer_chart')
    # Perform any required calculations or data processing

    # Pass the data to the template
    context = {
        'similarity_score': similarity_score,
        'speedometer_chart': speedometer_chart,
        'doc1_summary': doc1_summary,
        'doc2_summary': doc2_summary,

    }

    return render(request, 'show_report.html', context)


def doc_view(request):
    # Handle the request and return the response
    # Example:
    return render(request, 'similarity_app/doc.html')


def upload_view(request):
    # Handle the request and return the response
    # Example:
    return redirect('similarity_app/')
